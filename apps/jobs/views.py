from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from datetime import datetime

from docx import Document
from io import BytesIO
from reportlab.pdfgen import canvas


from apps.users.models import Settings
from apps.jobs.models import Jobs, CV, ReadyCV


def index(request):
    jobs = Jobs.objects.all()
    setting = Settings.objects.latest('id')
    try:
        cv = CV.objects.get(user=request.user)
    except:
        pass

    try:
        cv2 = ReadyCV.objects.get(user=request.user)
    except:
        pass
    return render(request, "jobs/company_listing.html", locals())

@login_required(login_url='login')
def jobs_detail(request, id):
    setting = Settings.objects.latest('id')
    job = Jobs.objects.get(id=id)
    jobs_category = Jobs.objects.filter(category=job.category).order_by('category')

    return render(request, 'jobs/listing_single.html', locals())


def cv(request):
    setting = Settings.objects.latest('id')
    return render(request, "jobs/cv.html", locals())

def cv_add(request):
    setting = Settings.objects.latest('id')

    if request.method == "POST":
        avatar = request.FILES.get('avatar')
        phone_number = request.POST.get('phone_number')
        address = request.POST.get('address')
        citizenship = request.POST.get('citizenship')
        birthday = request.POST.get('birthday')
        gender = request.POST.get('gender')
        children= request.POST.get('children', False)
        child_name = request.POST.get('child_name')
        child_birthday = request.POST.get('child_birthday')
        source = request.POST.get('source')
        reason = request.POST.get('reason')
        interested = request.POST.get('interested', False)
        interested_position = request.POST.get('interested_position')
        candidate = request.POST.get('candidate', False)
        test = request.POST.get('test', False)
        test_year= request.POST.get('test_year')
        test_score = request.POST.get('test_score')
        exam = request.POST.get('exam', False)
        exam_year = request.POST.get('exam_year')
        exam_score = request.POST.get('exam_score')
        education = request.POST.get('education', False)
        education_info = request.POST.get('education_info')
        qualification = request.POST.get('qualification', False)
        qualification_name = request.POST.get('qualification_name')
        qualification_country = request.POST.get('qualification_country')
        qualification_city = request.POST.get('qualification_city')
        qualification_start = request.POST.get('qualification_start')
        qualification_end = request.POST.get('qualification_end')
        organisation_name = request.POST.get('organisation_name')
        organisation_address = request.POST.get('organisation_address')
        job_title = request.POST.get('job_title')
        last_job_start = request.POST.get('last_job_start')
        last_job_end = request.POST.get('last_job_end')
        job_title_duty = request.POST.get('job_title_duty')
        base_moves = request.POST.get('base_moves')
        dismissal_reason = request.POST.get('dismissal_reason')
        advisor = request.POST.get('advisor')
        advisor_organisation = request.POST.get('advisor_organisation')
        advisor_start = request.POST.get('advisor_start')
        advisor_end = request.POST.get('advisor_end')
        advisor_address = request.POST.get('advisor_address')
        advisor_phone = request.POST.get('advisor_phone')
        advisor_socials = request.POST.get('advisor_socials')
        advisor_job = request.POST.get('advisor_job')
        internship_salary = request.POST.get('internship_salary')
        job_salary = request.POST.get('job_salary')
        language = request.POST.get('language')
        language_knowledge = request.POST.get('language_knowledge')
        language_years = request.POST.get('language_years')
        docs = request.POST.get('docs', False)
        tables = request.POST.get('tables', False)
        presentation = request.POST.get('presentation', False)
        prezi = request.POST.get('prezi', False)
        touch_typing = request.POST.get('touch_typing', False)
        abroad_country = request.POST.get('abroad_country')
        abroad_start = request.POST.get('abroad_start')
        abroad_end = request.POST.get('abroad_end')
        abroad_reason = request.POST.get('abroad_reason')
        good_skills = request.POST.get('good_skills')
        linkedin = request.POST.get('linkedin')
        facebook = request.POST.get('facebook')
        twitter = request.POST.get('twitter')
        university = request.POST.get('university')
        faculty = request.POST.get('faculty')
        edu_start = request.POST.get('edu_start')
        edu_end = request.POST.get('edu_end')
        edu_country = request.POST.get('edu_country')
        edu_city = request.POST.get('edu_city')
        plan = request.POST.get('plan')
        hobby = request.POST.get('hobby')
        who = request.POST.get('who')
        top_qualities = request.POST.get('top_qualities')
        worst_qualities = request.POST.get('worst_qualities')
        acq = request.POST.get('acq')
        acq_org = request.POST.get('acq_org')
        acq_job = request.POST.get('acq_job')
        acq_start = request.POST.get('acq_start')
        acq_end = request.POST.get('acq_end')
        acq_address = request.POST.get('acq_address')
        acq_phone = request.POST.get('acq_phone')
        acq_email = request.POST.get('acq_email')
        acq_socials = request.POST.get('acq_socials')
        achievement = request.POST.get('achievement')
        knowledge = request.POST.get('knowledge')
        historical_person = request.POST.get('hostorical_person')
        laptop = request.POST.get('laptop', False)
        change = request.POST.get('change')
        
        if not test_year or test_score or exam_year or exam_score or education_info or interested or candidate or test or exam or education or qualification:
            test_year = None 
            test_score = None 
            exam_year = None
            exam_score = None
            education_info = None
            interested = False
            candidate = False
            test = False
            exam = False
            education = False
            qualification= False
        
        cv = CV.objects.create( 
            user = request.user,
            avatar = avatar,
            phone_number = phone_number,
            address = address,
            citizenship = citizenship,
            birthday = birthday,
            gender = gender,
            children = children,
            child_name = child_name,
            child_birthday = child_birthday,
            source = source,
            reason = reason,
            interested = interested,
            interested_position = interested_position,
            candidate = candidate,
            republic_test = test,
            test_year = test_year,
            test_score = test_score,
            ielts_toefl = exam,
            exam_year = exam_year,
            exam_score = exam_score,
            education = education,
            education_info = education_info,
            qualification = qualification,
            qualification_name = qualification_name,
            qualification_country = qualification_country,
            qualification_city = qualification_city,
            qualification_start = qualification_start,
            qualification_end = qualification_end,
            organisation_name = organisation_name,
            organisation_address = organisation_address,
            job_title = job_title,
            last_job_start = last_job_start,
            last_job_end = last_job_end,
            job_title_duty = job_title_duty,
            base_moves = base_moves,
            dismissal_reason = dismissal_reason,
            advisor = advisor,
            advisor_organisation = advisor_organisation,
            advisor_start = advisor_start,
            advisor_end = advisor_end,
            advisor_address = advisor_address,
            advisor_phone = advisor_phone,
            advisor_socials = advisor_socials,
            advisor_job = advisor_job,
            internship_salary = internship_salary,
            job_salary = job_salary,
            language = language,
            language_knowledge = language_knowledge,
            language_years = language_years,
            google_docs = docs,
            google_tables = tables,
            google_presentation = presentation,
            prezi = prezi,
            touch_typing = touch_typing,
            abroad_country = abroad_country,
            abroad_start = abroad_start,
            abroad_end = abroad_end,
            abroad_reason = abroad_reason,
            good_skills = good_skills,
            linkedin = linkedin,
            facebook = facebook,
            twitter = twitter,
            university = university,
            faculty = faculty,
            edu_start = edu_start,
            edu_end = edu_end,
            edu_country = edu_country,
            edu_city = edu_city,
            plan = plan,
            hobby = hobby,
            who = who,
            top_qualities = top_qualities,
            worst_qualities = worst_qualities,
            acq = acq,
            acq_org = acq_org,
            acq_title = acq_job,
            acq_start = acq_start,
            acq_end = acq_end,
            acq_address = acq_address,
            acq_phone = acq_phone,
            acq_email = acq_email,
            acq_socials = acq_socials,
            achievement = achievement,
            knowledge = knowledge,
            historical_person = historical_person,
            laptop = laptop,
            change = change
        )
        cv.save()
        return redirect("/")
    return render(request, "jobs/add_cv.html", locals())


def cv_download(request):
    setting = Settings.objects.latest('id')

    if request.method == "POST":
        avatar = request.FILES.get('avatar')
        cv_file  = request.FILES.get('cv_file')
        checkbox = request.POST.get('shipping-option')

        if not checkbox:
            messages.error(request, 'Пожалуйста, согласитесь с условиями использования сайта.')
            return redirect('cv_download')

        if not cv_file.name.endswith(('.pdf', '.docx')):
            messages.error(request, 'Пожалуйста, загрузите резюме в формате PDF или Word.')
            return redirect('cv_download')
        
        ready_cv = ReadyCV.objects.create(user=request.user, avatar=avatar, cv_file=cv_file)
        ready_cv.save()

        if cv_file.name.endswith('.docx'):
            doc = Document()
            doc.add_heading('Resume', 0)
            doc.add_paragraph(f"User: {request.user}")
            wordeam = BytesIO()
            doc.save(wordeam)
            wordeam.seek(0)

        elif cv_file.name.endswith('.pdf'):
            pdfeam = BytesIO()
            pdf_canvas = canvas.Canvas(pdfeam)
            pdf_canvas.drawString(100, 800, f"Name: {request.user}")
            pdf_canvas.save()
            pdfeam.seek(0)

        return redirect('/')

    return render(request, "jobs/cv_download.html", locals())